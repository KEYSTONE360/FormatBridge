from __future__ import annotations

import html
import json
import shutil
import subprocess
import tempfile
from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path

from .engines import EngineManager
from .registry import (
    AUDIO_EXTENSIONS,
    HWP_EXTENSIONS,
    IMAGE_EXTENSIONS,
    PDF_EXTENSIONS,
    SHEET_EXTENSIONS,
    SLIDE_EXTENSIONS,
    TEXT_EXTENSIONS,
    VECTOR_EXTENSIONS,
    VIDEO_EXTENSIONS,
    WORD_EXTENSIONS,
    category,
    full_suffix,
)

StatusCallback = Callable[[str], None]


@dataclass
class ConvertOptions:
    quality: int = 88
    dpi: int = 160
    overwrite: bool = False
    pdf_quality: str = "전자책"
    background: str = "#FFFFFF"


def unique_output(output_dir: Path, stem: str, target_ext: str, overwrite: bool = False) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    target_ext = target_ext if target_ext.startswith(".") else f".{target_ext}"
    candidate = output_dir / f"{stem}{target_ext}"
    if overwrite or not candidate.exists():
        return candidate
    index = 2
    while True:
        candidate = output_dir / f"{stem}_{index}{target_ext}"
        if not candidate.exists():
            return candidate
        index += 1


def _flatten_for_jpeg(image, background: str):
    from PIL import Image, ImageColor

    if image.mode in {"RGBA", "LA"} or (image.mode == "P" and "transparency" in image.info):
        rgba = image.convert("RGBA")
        canvas = Image.new("RGB", rgba.size, ImageColor.getrgb(background))
        canvas.paste(rgba, mask=rgba.getchannel("A"))
        return canvas
    return image.convert("RGB")


class Converter:
    def __init__(self, engines: EngineManager | None = None):
        self.engines = engines or EngineManager()

    def convert(
        self,
        source: Path,
        target_ext: str,
        output_dir: Path,
        options: ConvertOptions | None = None,
        status: StatusCallback | None = None,
    ) -> list[Path]:
        options = options or ConvertOptions()
        source = Path(source)
        target_ext = target_ext.lower()
        if not target_ext.startswith("."):
            target_ext = f".{target_ext}"
        if not source.is_file():
            raise FileNotFoundError(f"입력 파일을 찾을 수 없습니다: {source}")

        source_ext = full_suffix(source)
        source_kind = category(source)
        if status:
            status(f"{source.name}: {source_ext} → {target_ext}")

        if source_ext in IMAGE_EXTENSIONS:
            return [self._convert_image(source, target_ext, output_dir, options)]
        if source_ext in VECTOR_EXTENSIONS:
            return [self._convert_vector(source, target_ext, output_dir, options)]
        if source_ext in PDF_EXTENSIONS:
            return self._convert_pdf(source, target_ext, output_dir, options)
        if source_ext in HWP_EXTENSIONS:
            output = unique_output(output_dir, source.stem, target_ext, options.overwrite)
            self.engines.office_convert("hwp", source, output)
            return [output]
        if source_ext in SHEET_EXTENSIONS:
            output = unique_output(output_dir, source.stem, target_ext, options.overwrite)
            self._office_or_libre("excel", source, output)
            return [output]
        if source_ext in SLIDE_EXTENSIONS:
            output = unique_output(output_dir, source.stem, target_ext, options.overwrite)
            self._office_or_libre("powerpoint", source, output)
            return [output]
        if source_ext in WORD_EXTENSIONS or source_ext in TEXT_EXTENSIONS:
            if target_ext == ".pdf" and source_ext in TEXT_EXTENSIONS:
                output = unique_output(output_dir, source.stem, target_ext, options.overwrite)
                self._text_to_pdf(source, output)
                return [output]
            output = unique_output(output_dir, source.stem, target_ext, options.overwrite)
            self._office_or_libre("word", source, output)
            return [output]
        if source_ext in AUDIO_EXTENSIONS or source_ext in VIDEO_EXTENSIONS:
            output = unique_output(output_dir, source.stem, target_ext, options.overwrite)
            self.engines.ffmpeg_convert(source, output, options.quality)
            return [output]
        raise ValueError(f"현재 설치된 엔진으로 {source_kind}({source_ext}) 변환을 처리할 수 없습니다.")

    def _office_or_libre(self, engine: str, source: Path, output: Path) -> None:
        if self.engines.ms_office:
            self.engines.office_convert(engine, source, output)
            return
        if self.engines.soffice:
            with tempfile.TemporaryDirectory(prefix="formatbridge_lo_") as temp:
                temp_path = Path(temp)
                cmd = [
                    self.engines.soffice,
                    "--headless",
                    "--convert-to",
                    output.suffix.lstrip("."),
                    "--outdir",
                    str(temp_path),
                    str(source),
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=300, check=False)
                candidates = list(temp_path.glob(f"{source.stem}.*"))
                if result.returncode != 0 or not candidates:
                    raise RuntimeError((result.stderr or result.stdout or "LibreOffice 변환 실패").strip())
                shutil.move(str(candidates[0]), output)
            return
        raise RuntimeError("이 문서 변환에는 Microsoft Office 또는 LibreOffice가 필요합니다.")

    def _convert_image(self, source: Path, target_ext: str, output_dir: Path, options: ConvertOptions) -> Path:
        from PIL import Image, ImageSequence

        output = unique_output(output_dir, source.stem, target_ext, options.overwrite)
        with Image.open(source) as image:
            frames = [frame.copy() for frame in ImageSequence.Iterator(image)]
            if not frames:
                frames = [image.copy()]
        target = target_ext.lower()
        if target == ".pdf":
            prepared = [_flatten_for_jpeg(frame, options.background) for frame in frames]
            prepared[0].save(output, "PDF", save_all=True, append_images=prepared[1:], resolution=options.dpi)
            return output

        format_map = {
            ".jpg": "JPEG",
            ".jpeg": "JPEG",
            ".png": "PNG",
            ".webp": "WEBP",
            ".avif": "AVIF",
            ".jp2": "JPEG2000",
            ".j2k": "JPEG2000",
            ".bmp": "BMP",
            ".gif": "GIF",
            ".tif": "TIFF",
            ".tiff": "TIFF",
            ".ico": "ICO",
        }
        if target not in format_map:
            raise ValueError(f"지원하지 않는 이미지 대상 형식입니다: {target}")
        fmt = format_map[target]
        prepared = frames
        kwargs: dict = {}
        if fmt in {"JPEG", "JPEG2000", "BMP"}:
            prepared = [_flatten_for_jpeg(frame, options.background) for frame in frames]
        if fmt in {"JPEG", "WEBP", "AVIF"}:
            kwargs["quality"] = max(1, min(100, options.quality))
        if fmt == "JPEG":
            kwargs.update(optimize=True, progressive=True)
        if fmt in {"GIF", "TIFF", "WEBP"} and len(prepared) > 1:
            kwargs.update(save_all=True, append_images=prepared[1:])
        prepared[0].save(output, fmt, **kwargs)
        return output

    def _convert_vector(self, source: Path, target_ext: str, output_dir: Path, options: ConvertOptions) -> Path:
        output = unique_output(output_dir, source.stem, target_ext, options.overwrite)
        if source.suffix.lower() == ".svg":
            try:
                import pymupdf as fitz
                from PIL import Image
            except ImportError as exc:
                raise RuntimeError("SVG 변환에는 MuPDF와 Pillow 엔진이 필요합니다.") from exc
            with fitz.open(stream=source.read_bytes(), filetype="svg") as document:
                if target_ext == ".pdf":
                    output.write_bytes(document.convert_to_pdf())
                    return output
                scale = options.dpi / 72.0
                pixmap = document[0].get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=True)
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
                    temp_path = Path(temp.name)
                try:
                    pixmap.save(str(temp_path))
                    with Image.open(temp_path) as image:
                        prepared = (
                            _flatten_for_jpeg(image, options.background)
                            if target_ext in {".jpg", ".jpeg", ".jp2", ".j2k", ".bmp"}
                            else image.copy()
                        )
                        fmt = {
                            ".jpg": "JPEG",
                            ".jpeg": "JPEG",
                            ".png": "PNG",
                            ".webp": "WEBP",
                            ".avif": "AVIF",
                            ".jp2": "JPEG2000",
                            ".j2k": "JPEG2000",
                            ".bmp": "BMP",
                            ".gif": "GIF",
                            ".tif": "TIFF",
                            ".tiff": "TIFF",
                            ".ico": "ICO",
                        }.get(target_ext)
                        if not fmt:
                            raise ValueError(f"SVG 대상 형식을 지원하지 않습니다: {target_ext}")
                        kwargs = {"quality": options.quality} if fmt in {"JPEG", "WEBP", "AVIF"} else {}
                        prepared.save(output, fmt, **kwargs)
                    return output
                finally:
                    temp_path.unlink(missing_ok=True)
        if self.engines.ghostscript and target_ext == ".pdf":
            result = subprocess.run(
                [
                    self.engines.ghostscript,
                    "-dBATCH",
                    "-dNOPAUSE",
                    "-sDEVICE=pdfwrite",
                    f"-sOutputFile={output}",
                    str(source),
                ],
                capture_output=True,
                text=True,
                timeout=300,
                check=False,
            )
            if result.returncode == 0 and output.exists():
                return output
        raise ValueError(f"{source.suffix} → {target_ext} 변환 엔진이 없습니다.")

    def _convert_pdf(self, source: Path, target_ext: str, output_dir: Path, options: ConvertOptions) -> list[Path]:
        if target_ext == ".pdf":
            output = unique_output(output_dir, f"{source.stem}_optimized", ".pdf", options.overwrite)
            self.engines.optimize_pdf(source, output, options.pdf_quality)
            return [output]
        if target_ext == ".docx":
            output = unique_output(output_dir, source.stem, target_ext, options.overwrite)
            self._pdf_to_docx(source, output, options)
            return [output]
        if target_ext == ".txt":
            try:
                import pymupdf as fitz
            except ImportError as exc:
                raise RuntimeError("PDF 텍스트 추출에는 PyMuPDF 엔진이 필요합니다.") from exc
            output = unique_output(output_dir, source.stem, ".txt", options.overwrite)
            with fitz.open(source) as doc, output.open("w", encoding="utf-8") as handle:
                for page_index, page in enumerate(doc, 1):
                    if page_index > 1:
                        handle.write(f"\n\n--- {page_index}쪽 ---\n\n")
                    handle.write(page.get_text())
            return [output]
        if target_ext not in IMAGE_EXTENSIONS:
            raise ValueError(f"PDF → {target_ext} 변환을 지원하지 않습니다.")
        try:
            import pymupdf as fitz
            from PIL import Image
        except ImportError as exc:
            raise RuntimeError("PDF 이미지 변환에는 PyMuPDF와 Pillow가 필요합니다.") from exc
        outputs: list[Path] = []
        scale = options.dpi / 72.0
        with fitz.open(source) as doc:
            for index, page in enumerate(doc, 1):
                pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=True)
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
                    temp_path = Path(temp.name)
                try:
                    pix.save(str(temp_path))
                    stem = f"{source.stem}_p{index:03d}"
                    output = unique_output(output_dir, stem, target_ext, options.overwrite)
                    with Image.open(temp_path) as image:
                        prepared = (
                            _flatten_for_jpeg(image, options.background)
                            if target_ext in {".jpg", ".jpeg"}
                            else image.copy()
                        )
                        prepared.save(
                            output,
                            format={
                                ".jpg": "JPEG",
                                ".jpeg": "JPEG",
                                ".png": "PNG",
                                ".webp": "WEBP",
                                ".tif": "TIFF",
                                ".tiff": "TIFF",
                            }.get(target_ext, target_ext[1:].upper()),
                            quality=options.quality,
                        )
                    outputs.append(output)
                finally:
                    temp_path.unlink(missing_ok=True)
        return outputs

    def _pdf_to_docx(self, source: Path, output: Path, options: ConvertOptions) -> None:
        """Create an unattended DOCX: editable text when present, page image otherwise."""
        try:
            import pymupdf as fitz
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError as exc:
            raise RuntimeError("PDF→DOCX 변환에는 MuPDF와 python-docx가 필요합니다.") from exc
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        normal = document.styles["Normal"]
        normal.font.name = "Malgun Gothic"
        normal.font.size = Pt(10)
        usable_width = section.page_width - section.left_margin - section.right_margin
        with fitz.open(source) as pdf, tempfile.TemporaryDirectory(prefix="formatbridge_docx_") as temp:
            temp_dir = Path(temp)
            for page_index, page in enumerate(pdf):
                blocks = [block for block in page.get_text("blocks", sort=True) if str(block[4]).strip()]
                if blocks:
                    for block in blocks:
                        text = str(block[4]).strip()
                        paragraph = document.add_paragraph()
                        for line_index, line in enumerate(text.splitlines()):
                            if line_index:
                                paragraph.add_run().add_break()
                            paragraph.add_run(line)
                else:
                    scale = min(max(options.dpi, 120), 220) / 72.0
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
                    image_path = temp_dir / f"page_{page_index + 1:04d}.png"
                    pixmap.save(str(image_path))
                    document.add_picture(str(image_path), width=usable_width)
                if page_index + 1 < len(pdf):
                    document.add_page_break()
        document.save(output)

    def _text_to_pdf(self, source: Path, output: Path) -> None:
        try:
            from reportlab.lib.enums import TA_LEFT
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as exc:
            raise RuntimeError("텍스트 PDF 변환에는 ReportLab 엔진이 필요합니다.") from exc

        text = source.read_text(encoding="utf-8", errors="replace")
        if source.suffix.lower() == ".json":
            try:
                text = json.dumps(json.loads(text), ensure_ascii=False, indent=2)
            except Exception:
                pass
        font_name = "Helvetica"
        for font in (r"C:\Windows\Fonts\malgun.ttf", r"C:\Windows\Fonts\gulim.ttc"):
            if Path(font).exists():
                try:
                    pdfmetrics.registerFont(TTFont("KoreanFont", font))
                    font_name = "KoreanFont"
                    break
                except Exception:
                    continue
        styles = getSampleStyleSheet()
        style = styles["BodyText"]
        style.fontName = font_name
        style.fontSize = 10
        style.leading = 15
        style.alignment = TA_LEFT
        doc = SimpleDocTemplate(str(output), pagesize=A4, rightMargin=48, leftMargin=48, topMargin=48, bottomMargin=48)
        story = []
        for line in text.splitlines() or [""]:
            safe = html.escape(line).replace(" ", "&nbsp;") or "&nbsp;"
            story.extend([Paragraph(safe, style), Spacer(1, 3)])
        doc.build(story)

    def merge_to_pdf(
        self,
        sources: list[Path],
        output: Path,
        options: ConvertOptions | None = None,
        status: StatusCallback | None = None,
    ) -> Path:
        options = options or ConvertOptions()
        if not sources:
            raise ValueError("합칠 파일이 없습니다.")
        try:
            from pypdf import PdfReader, PdfWriter
        except ImportError as exc:
            raise RuntimeError("PDF 병합에는 pypdf 엔진이 필요합니다.") from exc
        output.parent.mkdir(parents=True, exist_ok=True)
        writer = PdfWriter()
        with tempfile.TemporaryDirectory(prefix="formatbridge_merge_") as temp:
            temp_dir = Path(temp)
            for index, source in enumerate(sources, 1):
                if status:
                    status(f"PDF 준비 {index}/{len(sources)}: {source.name}")
                if source.suffix.lower() == ".pdf":
                    pdf = source
                else:
                    converted = self.convert(source, ".pdf", temp_dir, options, status)
                    if len(converted) != 1:
                        raise RuntimeError(f"PDF 준비 결과가 올바르지 않습니다: {source.name}")
                    pdf = converted[0]
                reader = PdfReader(str(pdf))
                for page in reader.pages:
                    writer.add_page(page)
            with output.open("wb") as handle:
                writer.write(handle)
        if not output.exists() or output.stat().st_size == 0:
            raise RuntimeError("병합 PDF가 생성되지 않았습니다.")
        return output


def verify_output(path: Path, expected_ext: str | None = None) -> tuple[bool, str]:
    if not path.exists():
        return False, "파일 없음"
    if path.stat().st_size <= 0:
        return False, "빈 파일"
    if expected_ext and full_suffix(path) != expected_ext.lower():
        return False, f"확장자 불일치: {full_suffix(path)}"
    ext = full_suffix(path)
    try:
        if ext in IMAGE_EXTENSIONS:
            from PIL import Image

            with Image.open(path) as image:
                image.verify()
            return True, "이미지 디코딩 정상"
        if ext == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            if len(reader.pages) < 1:
                return False, "PDF 페이지 없음"
            return True, f"PDF {len(reader.pages)}쪽 정상"
        if ext in {".docx", ".xlsx", ".pptx", ".hwpx"}:
            import zipfile

            with zipfile.ZipFile(path) as archive:
                if archive.testzip() is not None:
                    return False, "OOXML 내부 손상"
            return True, "OOXML 구조 정상"
        if ext == ".hwp":
            # HWP 5.x is an OLE Compound File with this fixed 8-byte signature.
            with path.open("rb") as handle:
                signature = handle.read(8)
            if signature != bytes.fromhex("D0CF11E0A1B11AE1"):
                return False, "HWP OLE 시그니처 불일치"
            return True, "네이티브 HWP 시그니처 정상"
    except Exception as exc:
        return False, str(exc)
    return True, "크기·존재 검사 정상"
